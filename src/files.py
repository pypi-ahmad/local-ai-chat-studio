"""Parse uploaded files into text, and prepare images for vision models."""

from __future__ import annotations

import io
from dataclasses import dataclass

from loguru import logger

IMAGE_TYPES = {"png", "jpg", "jpeg", "webp", "gif", "bmp"}
EXCEL_TYPES = {"xlsx", "xlsm", "xls"}
DOC_TYPES = {
    "pdf", "txt", "md", "csv", "tsv", "docx", "doc", "json", "py", "log",
    "yaml", "yml", "toml", "html",
} | EXCEL_TYPES
ACCEPTED_TYPES = sorted(IMAGE_TYPES | DOC_TYPES)

MIME_BY_EXT = {
    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "webp": "image/webp", "gif": "image/gif", "bmp": "image/bmp",
}


@dataclass
class Attachment:
    name: str
    kind: str  # 'image' | 'document'
    text: str = ""  # extracted text (documents, or OCR result for images)
    image_bytes: bytes | None = None
    mime: str = ""  # for images, e.g. 'image/png'


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def parse_upload(filename: str, raw: bytes) -> Attachment:
    """Convert an uploaded file into an Attachment (text extracted for docs)."""
    ext = _ext(filename)
    if ext in IMAGE_TYPES:
        return Attachment(
            name=filename, kind="image", image_bytes=raw,
            mime=MIME_BY_EXT.get(ext, "image/png"),
        )
    if ext == "pdf":
        return Attachment(name=filename, kind="document", text=_parse_pdf(raw))
    if ext == "docx":
        return Attachment(name=filename, kind="document", text=_parse_docx(raw))
    if ext == "doc":
        return Attachment(name=filename, kind="document", text=_parse_legacy_doc(raw))
    if ext in EXCEL_TYPES:
        return Attachment(name=filename, kind="document", text=_parse_excel(raw, ext))
    # Everything else: treat as utf-8 text
    try:
        text = raw.decode("utf-8", errors="replace")
    except Exception:
        text = ""
    return Attachment(name=filename, kind="document", text=text)


def _parse_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("PDF page {} extraction failed: {}", i, exc)
    return "\n\n".join(pages).strip()


def _parse_docx(raw: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def _parse_excel(raw: bytes, ext: str) -> str:
    import pandas as pd

    engine = "xlrd" if ext == "xls" else "openpyxl"
    sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None, engine=engine)
    parts = []
    for name, df in sheets.items():
        parts.append(f"## Sheet: {name}\n{df.to_csv(index=False)}")
    return "\n\n".join(parts).strip()


def _parse_legacy_doc(raw: bytes) -> str:
    """Old binary .doc — try antiword, then LibreOffice conversion."""
    import shutil
    import subprocess
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        path = f"{tmp}/file.doc"
        with open(path, "wb") as f:
            f.write(raw)
        if shutil.which("antiword"):
            try:
                out = subprocess.run(
                    ["antiword", path], capture_output=True, text=True, timeout=30
                )
                if out.returncode == 0 and out.stdout.strip():
                    return out.stdout.strip()
            except Exception as exc:
                logger.warning("antiword failed: {}", exc)
        if shutil.which("soffice"):
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt", "--outdir", tmp, path],
                    capture_output=True, timeout=60,
                )
                txt = f"{tmp}/file.txt"
                from pathlib import Path
                if Path(txt).exists():
                    return Path(txt).read_text(errors="replace").strip()
            except Exception as exc:
                logger.warning("soffice conversion failed: {}", exc)
    logger.warning("No .doc converter available (install antiword or libreoffice)")
    return ""


def chunk_text(text: str, chunk_chars: int, overlap_chars: int) -> list[str]:
    """Split text into overlapping chunks on paragraph boundaries where possible."""
    if len(text) <= chunk_chars:
        return [text] if text.strip() else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_chars, len(text))
        if end < len(text):
            # Prefer to break at a paragraph or sentence boundary near the end
            window = text[start:end]
            for sep in ("\n\n", "\n", ". "):
                cut = window.rfind(sep)
                if cut > chunk_chars // 2:
                    end = start + cut + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap_chars, start + 1)
    return chunks
